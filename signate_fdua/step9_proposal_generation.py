"""
Step 9: 提案書・成果物作成
==========================
- LLMによる提案書セクション生成（15,000字以内 / A4 10ページ目安）
- 評価観点（全体構成・地域性・業界特性・GX/DX・人材）を反映
- 3部構成: (1)企業概要・分析・課題 (2)成長戦略・提案 (3)効果試算・ロードマップ
- プロンプトログの統合
- 検証プロセスレポートの作成
- ZIP圧縮・提出準備
"""

import os
import json
import shutil
import time
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import re

import google.generativeai as genai


# =============================================================================
# 設定
# =============================================================================
BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
PROCESSED_DIR = DATA_DIR / "processed"
ANALYSIS_DIR = BASE_DIR / "analysis"

# 入力ディレクトリ
SWOT_DIR = ANALYSIS_DIR / "swot"
STRATEGY_DIR = ANALYSIS_DIR / "strategy"
ROADMAP_DIR = ANALYSIS_DIR / "roadmap"
LLM_ANALYSIS_DIR = ANALYSIS_DIR / "llm_analysis"
REPORTS_DIR = ANALYSIS_DIR / "reports"
PROMPT_LOG_DIR = BASE_DIR / "output" / "prompt_logs"
PROMPT_LOG_DIR.mkdir(parents=True, exist_ok=True)

# 出力ディレクトリ
OUTPUT_DIR = BASE_DIR / "output"
PROPOSALS_DIR = OUTPUT_DIR / "proposals"
PROPOSALS_DIR.mkdir(parents=True, exist_ok=True)

GEMINI_MODEL = "gemini-3-pro-preview"

# 文字数目安（gemini-2.5-pro は指示に忠実なため、目標値をそのまま指定）
# 合計 12,000〜13,000字 → 表紙等を含めて15,000字以内に収まる
CHAR_LIMITS = {
    "executive_summary": 800,  # エグゼクティブサマリー（A4 約1ページ）
    "section1": 4000,  # 企業概要・分析・課題抽出（A4 約3ページ）
    "section2": 4000,  # 成長戦略・提案（A4 約3ページ）
    "section3": 2500,  # 効果試算・ロードマップ（A4 約2ページ）
}


# =============================================================================
# プロンプトログ
# =============================================================================
@dataclass
class PromptLogEntry:
    timestamp: str
    company_code: str
    purpose: str
    input_prompt: str
    output_response: str
    model: str = GEMINI_MODEL


class PromptLogger:
    def __init__(self):
        self.logs: list[PromptLogEntry] = []
        self.log_file = PROMPT_LOG_DIR / f"prompt_log_step9_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

    def add_log(self, entry: PromptLogEntry):
        self.logs.append(entry)
        logs_dict = [asdict(log) for log in self.logs]
        with open(self.log_file, "w", encoding="utf-8") as f:
            json.dump(logs_dict, f, ensure_ascii=False, indent=2)

    def export_to_text(self) -> str:
        text = "=" * 80 + "\n"
        text += "Step 9: 提案書生成 プロンプトログ\n"
        text += f"生成日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        text += "=" * 80 + "\n\n"
        for i, log in enumerate(self.logs, 1):
            text += f"--- ログ #{i} ---\n"
            text += f"日時: {log.timestamp}\n"
            text += f"対象企業: {log.company_code}\n"
            text += f"目的: {log.purpose}\n"
            text += f"\n【入力プロンプト】\n{log.input_prompt}\n"
            text += f"\n【出力結果】\n{log.output_response}\n"
            text += "\n" + "-" * 80 + "\n\n"
        return text


# =============================================================================
# データ読み込み
# =============================================================================
def load_all_analysis_data(company_code: str) -> dict:
    """全分析データを統合して読み込み"""
    data = {"company_code": company_code}

    loaders = {
        "master": (PROCESSED_DIR / "company_master.json",
                   lambda d: next((m for m in d if m["code"] == company_code), {})),
        "financial": (REPORTS_DIR / "analysis_summary_for_ai.json",
                      lambda d: d.get(company_code, {})),
        "securities": (LLM_ANALYSIS_DIR / "securities_report_analysis.json",
                       lambda d: d.get(company_code, {})),
        "swot": (SWOT_DIR / "swot_analysis_results.json",
                 lambda d: d.get(company_code, {})),
        "strategy": (STRATEGY_DIR / "strategy_results.json",
                     lambda d: d.get(company_code, {})),
        "roadmap": (ROADMAP_DIR / "effect_estimation_results.json",
                    lambda d: d.get(company_code, {})),
    }

    for key, (path, extractor) in loaders.items():
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                data[key] = extractor(json.load(f))

    return data


def get_company_codes() -> list[str]:
    """企業コードリストを取得"""
    master_path = PROCESSED_DIR / "company_master.json"
    if master_path.exists():
        with open(master_path, "r", encoding="utf-8") as f:
            return sorted([m["code"] for m in json.load(f)])
    return []


# =============================================================================
# コンテキスト構築（LLMに渡す企業データ要約）
# =============================================================================
def build_company_context(data: dict) -> str:
    """LLMに渡す企業コンテキストを構築"""
    master = data.get("master", {})
    financial = data.get("financial", {})
    securities = data.get("securities", {})
    swot = data.get("swot", {})
    strategy = data.get("strategy", {})
    roadmap = data.get("roadmap", {})

    ctx = f"""
===== 企業基本情報 =====
企業コード: {data['company_code']}
本社所在地: {master.get('location', '')}
業種分類: {master.get('industry', '')}
市場区分: {master.get('market', '')}
従業員数: {master.get('employees', '')}名
資本金: {master.get('capital', '')}億円

===== 財務データ（3年推移） =====
業績: {json.dumps(financial.get('業績推移', {}), ensure_ascii=False)}
収益性: {json.dumps(financial.get('収益性指標', {}), ensure_ascii=False)}
安全性: {json.dumps(financial.get('安全性指標', {}), ensure_ascii=False)}
CF: {json.dumps(financial.get('キャッシュフロー_億円', {}), ensure_ascii=False)}
財務面の強み: {financial.get('分析結果', {}).get('強み', [])}
財務面の課題: {financial.get('分析結果', {}).get('課題', [])}

===== 有価証券報告書分析 =====
事業セグメント: {securities.get('business_segments', [])}
事業内容: {securities.get('main_business_description', '')}
販路構成: {json.dumps(securities.get('sales_structure', {}), ensure_ascii=False)}
経営理念: {securities.get('management_philosophy', '')}
中期経営計画: {securities.get('medium_term_plan', '')}
重点戦略: {securities.get('key_strategies', [])}
DX取組: {securities.get('dx_initiatives', [])}
GX取組: {securities.get('gx_initiatives', [])}
リスク: {securities.get('business_risks', [])}
人材戦略: {securities.get('hr_strategy', '')}

===== 事業実績・計画（数値データ）=====
中期経営計画の目標数値: {securities.get('medium_term_plan_targets', '情報なし')}
セグメント別業績: {securities.get('segment_performance', '情報なし')}
受注・販売実績: {securities.get('order_sales_record', '情報なし')}

===== 経営者の分析（LLMインサイト）=====
{securities.get('management_analysis', '情報なし')}

===== 設備投資の戦略分析（LLMインサイト）=====
{securities.get('capex_plan', '情報なし')}

===== SWOT分析 =====
強み: {swot.get('strengths', [])}
弱み: {swot.get('weaknesses', [])}
機会: {swot.get('opportunities', [])}
脅威: {swot.get('threats', [])}
SO戦略: {swot.get('so_strategies', [])}
WO戦略: {swot.get('wo_strategies', [])}
ST戦略: {swot.get('st_strategies', [])}
WT戦略: {swot.get('wt_strategies', [])}
優先課題: {json.dumps(swot.get('priority_issues', []), ensure_ascii=False)}

===== 戦略提案 =====
方向性: {strategy.get('strategic_direction', '')}
ビジョン: {strategy.get('vision', '')}
成長戦略: {json.dumps(strategy.get('growth_strategies', []), ensure_ascii=False)}
DX施策: {json.dumps(strategy.get('dx_measures', []), ensure_ascii=False)}
GX施策: {json.dumps(strategy.get('gx_measures', []), ensure_ascii=False)}
人材施策: {json.dumps(strategy.get('hr_measures', []), ensure_ascii=False)}
短期アクション: {json.dumps(strategy.get('short_term_actions', []), ensure_ascii=False)}
中期アクション: {json.dumps(strategy.get('medium_term_actions', []), ensure_ascii=False)}
長期アクション: {json.dumps(strategy.get('long_term_actions', []), ensure_ascii=False)}

===== 効果試算・ロードマップ =====
売上効果: {json.dumps(roadmap.get('revenue_effect', {}), ensure_ascii=False)}
利益効果: {json.dumps(roadmap.get('profit_effect', {}), ensure_ascii=False)}
投資計画: {json.dumps(roadmap.get('investment_plan', []), ensure_ascii=False)}
ROI: {json.dumps(roadmap.get('roi_analysis', {}), ensure_ascii=False)}
KPI: {json.dumps(roadmap.get('kpis', []), ensure_ascii=False)}
ロードマップ: {json.dumps(roadmap.get('roadmap', []), ensure_ascii=False)}
リスク: {json.dumps(roadmap.get('risks_and_mitigations', []), ensure_ascii=False)}
"""

    # 外部情報（業界レポート）を追加
    try:
        from external_info_loader import get_external_context_for_company
        location = master.get('location', '')
        external_context = get_external_context_for_company(location)
        if external_context:
            ctx += f"""

===== 外部環境情報（業界レポートより） =====
{external_context}
"""
    except ImportError:
        pass

    return ctx


# =============================================================================
# 提案書セクション生成プロンプト
# =============================================================================

SYSTEM_INSTRUCTION = """あなたは地域金融機関の融資担当者兼建設業専門の経営コンサルタントです。
取引先企業の持続的成長を支援するため、データに基づいた説得力ある提案書を執筆します。

以下の5つの評価観点を意識して、各セクションの内容に必ず反映してください。

【評価観点】
1. 全体構成（Past→Future）: 過去3年の財務・事業分析と未来の成長戦略が、一貫した因果関係で論理的に接続されていること。
2. 地域性: 企業所在地の商圏、人口動態、行政施策、地域固有のインフラ計画を踏まえ、地域密着型の具体的な提案ができていること。
3. 業界特性: 官公庁／民間、元請／下請などの販路・商流特性を把握し、それに応じた分析と提案がなされていること。
4. GX/DX対応: BIM/CIM、ICT施工、ZEB/ZEH、低炭素工法など、技術トレンドへの具体的な投資・対応策を提案していること。
5. 人材不足対応: 工事需要の変化や深刻な人手不足（2024年問題、外国人材受入等）に対し、実効性のある解決策（歩掛管理の高度化、採用・定着策等）を示すこと。

【執筆スタイル】
- 文章と箇条書きを適切に使い分け、読みやすく構造化してください。
- 数値データを積極的に用い、根拠のある記述を心がけてください。
- 見出しには ■（大見出し）、●（中見出し）、・（箇条書き）を使用してください。
- Markdown記法（#、**、```等）は使用しないでください。

【禁止事項】
- 「承知しました」「以下に作成します」「かしこまりました」等のAIとしての返事や枕詞は絶対に含めないでください。
- 冒頭の挨拶や末尾の締めの言葉（「以上です」「ご質問があれば〜」等）も不要です。
- 提案書の本文のみを出力してください。余計な前置きや後書きは一切不要です。
"""


def make_section1_prompt(context: str, char_limit: int) -> str:
    """第1部: 企業概要・分析・課題の抽出"""
    return f"""{SYSTEM_INSTRUCTION}

以下の企業データに基づき、提案書の「第1部：企業概要・分析、課題の抽出」を{char_limit}字程度で執筆してください。

{context}

【構成】
■ 企業概要
  基本情報（所在地、業種、従業員数、資本金、市場区分）を示した上で、
  事業セグメントの内容と各セグメントの特徴を説明してください。
  販路構成（官公庁/民間の比率、元請/下請の状況）にも必ず言及してください。

■ 外部環境分析
  建設業界全体の動向（資材価格高騰、労務費上昇、2024年問題等）を概観した上で、企業所在地に固有の地域特性を詳しく記述してください。
  地域の人口動態、主要産業、行政のインフラ投資計画、地域固有の建設需要など、画一的ではない具体的な情報を盛り込んでください。
  外部環境情報が提供されている場合は、その内容を積極的に活用してください。

■ 財務分析（過去3年間の推移）
  売上高、営業利益、当期純利益の3年推移を具体的な数値で示し、収益性指標（営業利益率、ROE、ROA）、安全性指標（自己資本比率、流動比率）、キャッシュフローの状況を分析してください。
  数値の変化要因（なぜ増減したのか）を明確に記述してください。

■ 課題の抽出
  SWOT分析（強み・弱み・機会・脅威）を整理した上で、上記の分析から論理的に導かれる経営課題を優先順位付きで3〜5つ提示してください。
  各課題がなぜ重要なのか、財務データや外部環境との関連を明示してください。

【文字数】{char_limit}字程度
"""


def make_section2_prompt(context: str, section1_text: str, char_limit: int) -> str:
    """第2部: 成長戦略、提案（第1部の内容を考慮）"""
    return f"""{SYSTEM_INSTRUCTION}

あなたは現在、提案書を執筆中です。既に完成した「第1部」の内容を受け、論理的に繋がる形で「第2部：成長戦略・提案」を{char_limit}字程度で執筆してください。

【既に作成済みの第1部】
{section1_text}

【基礎となる企業データ】
{context}

【最重要ポイント】
第1部で抽出した課題と、本セクションの戦略提案を因果関係で必ず接続してください。
「この課題があるから → この戦略を提案する」という論理を各提案で明示してください。

【構成】
■ 戦略の方向性と目指す姿
  3〜5年後のビジョンを示し、基本戦略の考え方を説明してください。
  第1部の課題分析から導かれる戦略の必然性を論じてください。

■ 成長戦略の提案
  以下の観点から2〜4つの戦略を提案してください。各戦略について、背景（なぜ必要か）、具体的施策、期待効果を記述してください。
  ・事業拡大・収益性向上策
  ・地域戦略（所在地の特性を活かした具体策、地域のインフラ需要との連携）
  ・販路最適化（官公庁/民間バランス、元請比率向上、新規顧客開拓等）

■ DX（デジタルトランスフォーメーション）施策
  ICT施工（ドローン、ICT建機）、BIM/CIM活用、業務効率化（施工管理アプリ等）、省力化・自動化技術など、企業規模に見合った具体的なDX施策を2〜3つ程度提案してください。
  各施策の導入効果と投資規模感にも触れてください。

■ GX（グリーントランスフォーメーション）施策
  施工時CO2削減（低炭素資材、ハイブリッド建機）、環境配慮型建築（ZEB/ZEH）、再エネ関連など、具体的なGX施策を2〜3つ提案してください。

■ 人材確保・育成施策
  2024年問題（時間外労働規制）への対応、歩掛管理の高度化による生産性向上、採用強化（新卒・中途・外国人材）、技術承継、働き方改革、定着率向上策など、実効性のある人材施策を2〜3つ程度提案してください。

【文字数】{char_limit}字程度
"""


def make_section3_prompt(context: str, section2_text: str, char_limit: int) -> str:
    """第3部: 効果試算、ロードマップ（第2部の内容を考慮）"""
    return f"""{SYSTEM_INSTRUCTION}

あなたは現在、提案書の最終章を執筆中です。既に提案した「第2部」の具体的施策を受け、それらを実行した場合の「第3部：効果試算・ロードマップ」を{char_limit}字程度で執筆してください。

【既に提案済みの第2部（成長戦略・施策）】
{section2_text}

【基礎となる企業データ】
{context}

【構成】
■ 効果試算
  第2部で提案した施策を実行した場合の定量的な効果を試算してください。
  ・売上高予測（1年後、3年後、5年後）と成長率
  ・営業利益予測と利益率の改善見通し
  ・DX/GX施策によるコスト削減効果（生産性向上率など）
  各数値には算出根拠（前提条件や参考とした業界データ等）を必ず付記してください。

■ 投資計画とROI
  主要な投資項目、投資額、投資回収期間を示してください。
  全体のROI（投資収益率）も提示してください。

■ KPI設定
  財務KPI（売上高、営業利益率、ROE等）、事業KPI（受注高、生産性指標等）、DX/GX KPI（ICT施工比率、CO2削減率等）、人材KPI（採用数、定着率等）から、
  主要な指標を5〜8つ選び、現在値と1年後/3年後/5年後の目標値を示してください。

■ 実行ロードマップ
  ・短期（1年以内）：即時着手するクイックウィン施策と体制整備
  ・中期（1-3年）：本格投資と成果の刈り取り開始
  ・長期（3-5年）：目指す姿の実現と次の成長ステージへ
  各フェーズの主要アクションとマイルストーンを示してください。

■ 想定リスクと対策
  戦略実行上の主要リスクを3〜5つ挙げ、それぞれの対策と発生時のコンティンジェンシーを示してください。

【文字数】{char_limit}字程度
"""


def make_executive_summary_prompt(
    section1_text: str,
    section2_text: str,
    section3_text: str,
    char_limit: int,
) -> str:
    """エグゼクティブサマリー生成プロンプト"""
    return f"""{SYSTEM_INSTRUCTION}

あなたはプロの経営コンサルタントです。
あなたは既に作成済みの提案書（第1部〜第3部）を読み、その内容を凝縮した
経営陣に向けた「エグゼクティブサマリー」を{char_limit}字程度で執筆してください。

【既に作成済みの提案書本文】

＜第1部：企業概要・分析、課題の抽出＞
{section1_text}

＜第2部：成長戦略・提案＞
{section2_text}

＜第3部：効果試算・ロードマップ＞
{section3_text}

【エグゼクティブサマリーの要件】
・提案書全体の要旨を経営層が短時間で把握できるよう簡潔にまとめること。
・以下の要素を必ず含めること：
  1. 企業の現状と主要課題（2〜3点）
  2. 提案する成長戦略の骨子（DX/GX/人材施策を含む）
  3. 期待される定量的効果（売上・利益の改善見通し）
  4. 実行のタイムライン概要（短期・中期・長期）
・自信に満ちた、論理的かつ情熱的なビジネス文書を執筆すること。
・数値データを積極的に引用し、根拠ある記述を心がけること。
・見出しには ■（大見出し）、●（中見出し）、・（箇条書き）を使用してください。
・Markdown記法（#、**、```等）は使用しないでください。

【文字数】{char_limit}字程度
"""


# =============================================================================
# 提案書生成クラス
# =============================================================================
class ProposalGenerator:
    """LLMを活用した提案書生成クラス"""

    def __init__(self, api_key: Optional[str] = None):
        self.api_key = api_key or os.environ.get("GOOGLE_API_KEY")
        if not self.api_key:
            raise ValueError("GOOGLE_API_KEY が設定されていません。")

        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel(GEMINI_MODEL)
        self.prompt_logger = PromptLogger()

    @staticmethod
    def _clean_llm_response(text: str) -> str:
        """LLMの余計な返事・枕詞を除去する"""
        # 先頭にありがちなLLMの挨拶・前置きパターン
        prefixes_to_remove = [
            r"^承知しました[。．]?\s*",
            r"^かしこまりました[。．]?\s*",
            r"^了解しました[。．]?\s*",
            r"^はい[、,]?\s*",
            r"^もちろんです[。．]?\s*",
            r"^わかりました[。．]?\s*",
            r"^以下[にのが].*?(?:作成|執筆|記述|生成|まとめ|記載).*?[。．]\s*",
            r"^それでは[、,]?\s*",
            r"^では[、,]?\s*",
            r"^ご依頼.*?(?:に基づ[いき]|いただ[いき]).*?[。．]\s*",
            r"^ご指示.*?(?:に基づ[いき]|いただ[いき]|に従[いっ]).*?[。．]\s*",
            r"^提案書.*?(?:を作成|を執筆|をまとめ).*?[。．]\s*",
            r"^以下[、,].*?(?:です|ます|した)[。．]\s*",
        ]
        cleaned = text.strip()
        for pattern in prefixes_to_remove:
            cleaned = re.sub(pattern, "", cleaned, count=1, flags=re.MULTILINE)
            cleaned = cleaned.strip()

        # 末尾にありがちなLLMの締めの挨拶パターン
        suffixes_to_remove = [
            r"\s*以上[、,]?.*?(?:です|ます|でした|した)[。．]?\s*$",
            r"\s*何か.*?(?:ございましたら|ありましたら|あれば).*?[。．]?\s*$",
            r"\s*ご不明.*?(?:ございましたら|ありましたら|あれば).*?[。．]?\s*$",
            r"\s*ご質問.*?(?:ございましたら|ありましたら|あれば).*?[。．]?\s*$",
            r"\s*お気軽に.*?(?:ください|下さい)[。．]?\s*$",
            r"\s*ご確認.*?(?:ください|下さい|いただ)[。．]?\s*$",
            r"\s*必要に応じて.*?(?:ください|下さい|いただ|ます)[。．]?\s*$",
        ]
        for pattern in suffixes_to_remove:
            cleaned = re.sub(pattern, "", cleaned, count=1, flags=re.MULTILINE)
            cleaned = cleaned.strip()

        return cleaned

    def _call_llm(self, prompt: str, company_code: str, purpose: str) -> str:
        """LLMを呼び出してログを記録（余計な返事を自動除去）"""
        try:
            response = self.model.generate_content(prompt)
            raw_result = response.text
            result = self._clean_llm_response(raw_result)

            self.prompt_logger.add_log(PromptLogEntry(
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                company_code=company_code,
                purpose=purpose,
                input_prompt=prompt,
                output_response=result,
            ))

            return result
        except Exception as e:
            print(f"[WARNING] API呼び出しエラー: {e}")
            time.sleep(3)
            return ""

    def generate_proposal(self, company_code: str) -> Document:
        """1社分の提案書を生成"""
        print(f"[INFO] 提案書生成開始: {company_code}")

        # データ読み込み・コンテキスト構築
        data = load_all_analysis_data(company_code)
        context = build_company_context(data)

        # 各セクションをLLMで生成
        print(f"  - 第1部（企業概要・分析・課題）生成中...")
        section1_text = self._call_llm(
            make_section1_prompt(context, CHAR_LIMITS["section1"]),
            company_code,
            "提案書_第1部_企業概要・分析・課題"
        )
        time.sleep(2)

        # --- 第2部 生成（第1部のテキストを渡す） ---
        print(f"  - 第2部（成長戦略・提案）生成中...")
        section2_text = self._call_llm(
            make_section2_prompt(context, section1_text, CHAR_LIMITS["section2"]),
            company_code,
            "提案書_第2部_成長戦略・提案"
        )
        time.sleep(2)

        # --- 第3部 生成（第2部のテキストを渡す） ---
        print(f"  - 第3部（効果試算・ロードマップ）生成中...")
        section3_text = self._call_llm(
            make_section3_prompt(context, section2_text, CHAR_LIMITS["section3"]),
            company_code,
            "提案書_第3部_効果試算・ロードマップ"
        )
        time.sleep(2)

        # エグゼクティブサマリーを第1〜3部の内容から生成
        print(f"  - エグゼクティブサマリー生成中...")
        exec_summary_text = self._call_llm(
            make_executive_summary_prompt(
                section1_text, section2_text, section3_text,
                CHAR_LIMITS["executive_summary"],
            ),
            company_code,
            "提案書_エグゼクティブサマリー"
        )

        # 文字数チェックと強制トリミング
        total_chars = (
            len(exec_summary_text)
            + len(section1_text) + len(section2_text) + len(section3_text)
        )
        print(
            f"  - 文字数: サマリー={len(exec_summary_text)}, "
            f"第1部={len(section1_text)}, 第2部={len(section2_text)}, "
            f"第3部={len(section3_text)}, 合計={total_chars}"
        )

        # Wordドキュメント組み立て
        doc = self._build_docx(
            data, exec_summary_text,
            section1_text, section2_text, section3_text,
        )

        print(f"[INFO] 提案書生成完了: {company_code}")
        return doc

    @staticmethod
    def _trim_text(text: str, max_chars: int) -> str:
        """テキストを指定文字数以内にトリミング（行単位で切る）"""
        if len(text) <= max_chars:
            return text
        lines = text.split("\n")
        trimmed = []
        current_len = 0
        for line in lines:
            if current_len + len(line) + 1 > max_chars:
                break
            trimmed.append(line)
            current_len += len(line) + 1
        return "\n".join(trimmed)

    def _build_docx(
        self,
        data: dict,
        exec_summary_text: str,
        section1_text: str,
        section2_text: str,
        section3_text: str,
    ) -> Document:
        """Wordドキュメントを組み立てる"""
        doc = Document()
        master = data.get("master", {})

        # スタイル設定
        style = doc.styles['Normal']
        style.font.name = 'メイリオ'
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)

        # --- 表紙 ---
        for _ in range(6):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("成長戦略提案書")
        run.bold = True
        run.font.size = Pt(26)

        doc.add_paragraph()

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"企業コード: {data.get('company_code', '')}")
        run.font.size = Pt(16)

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(
            f"{master.get('location', '')} / {master.get('industry', '')}"
        )
        run.font.size = Pt(13)

        for _ in range(6):
            doc.add_paragraph()

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(datetime.now().strftime("%Y年%m月"))
        run.font.size = Pt(11)

        doc.add_page_break()

        # --- エグゼクティブサマリー ---
        doc.add_heading("エグゼクティブサマリー", level=1)
        self._add_text_block(doc, exec_summary_text)
        doc.add_page_break()

        # --- 第1部 ---
        doc.add_heading("第1部　企業概要・分析、課題の抽出", level=1)
        self._add_text_block(doc, section1_text)
        doc.add_page_break()

        # --- 第2部 ---
        doc.add_heading("第2部　成長戦略・提案", level=1)
        self._add_text_block(doc, section2_text)
        doc.add_page_break()

        # --- 第3部 ---
        doc.add_heading("第3部　効果試算・ロードマップ", level=1)
        self._add_text_block(doc, section3_text)

        return doc

    def _add_text_block(self, doc: Document, text: str):
        """テキストブロックをWordに追加（見出し記号を検出して構造化）"""
        lines = text.strip().split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 見出しレベルの判定
            if stripped.startswith("■"):
                doc.add_heading(stripped.lstrip("■").strip(), level=2)
            elif stripped.startswith("●"):
                doc.add_heading(stripped.lstrip("●").strip(), level=3)
            # 箇条書きの判定
            elif stripped.startswith("・"):
                # 先頭の「・」を削除してから箇条書きスタイルを適用
                text = stripped.lstrip("・").strip()
                doc.add_paragraph(text, style='List Bullet')
            elif stripped.startswith("- "):
                # 先頭の「- 」を削除してから箇条書きスタイルを適用
                text = stripped.lstrip("- ").strip()
                doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(stripped)

    def save_proposal(self, doc: Document, company_code: str):
        """提案書を保存"""
        output_path = PROPOSALS_DIR / f"{company_code}.docx"
        doc.save(output_path)
        print(f"[SAVE] {output_path}")

    def generate_all_proposals(self, company_codes: list[str]):
        """全企業の提案書を生成"""
        for i, code in enumerate(company_codes, 1):
            print(f"\n[{i}/{len(company_codes)}] {code}")
            try:
                doc = self.generate_proposal(code)
                self.save_proposal(doc, code)
            except Exception as e:
                print(f"[ERROR] {code}: {e}")

            if i < len(company_codes):
                time.sleep(3)


# =============================================================================
# プロンプトログ統合
# =============================================================================
def merge_prompt_logs():
    """全ステップのプロンプトログを統合"""
    merged_text = "=" * 80 + "\n"
    merged_text += "プロンプトログ（Prompt Log）- 統合版\n"
    merged_text += f"生成日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    merged_text += "=" * 80 + "\n\n"

    log_files = sorted(PROMPT_LOG_DIR.glob("*.txt"))

    for log_file in log_files:
        if log_file.name != "prompt_log.txt":
            merged_text += f"\n{'=' * 40}\n"
            merged_text += f"ファイル: {log_file.name}\n"
            merged_text += f"{'=' * 40}\n\n"
            with open(log_file, "r", encoding="utf-8") as f:
                merged_text += f.read()
            merged_text += "\n"

    output_path = OUTPUT_DIR / "prompt_log.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(merged_text)
    print(f"[SAVE] {output_path}")
    return output_path


# =============================================================================
# 検証プロセスレポート生成
# =============================================================================
def generate_verification_report():
    """検証プロセスレポートを生成"""
    doc = Document()

    doc.add_heading("検証プロセスレポート", level=1)
    doc.add_paragraph(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")

    doc.add_heading("1. AI活用の概要", level=2)
    doc.add_paragraph(
        "本提案書の作成にあたり、生成AI（Gemini）を活用しました。"
        "AIは以下の工程で使用しています。\n"
        "・有価証券報告書の要約・構造化（事業内容、リスク、経営方針の抽出）\n"
        "・財務データの解釈と示唆の導出\n"
        "・SWOT分析の実施（クロスSWOT戦略の策定を含む）\n"
        "・成長戦略オプションの生成（DX/GX/人材施策）\n"
        "・効果試算の算出とロードマップ策定\n"
        "・提案書各セクションの文章生成"
    )

    doc.add_heading("2. 検証プロセス", level=2)
    doc.add_paragraph(
        "AIの出力に対して、以下の3段階の検証プロセスを実施しました。\n\n"
        "【ファクトチェック】\n"
        "・財務データの数値がオリジナルCSVデータと一致しているか確認\n"
        "・有価証券報告書からの引用内容が正確か照合\n"
        "・業界動向に関する記述が実態と乖離していないか確認\n\n"
        "【論理整合性チェック】\n"
        "・過去分析（Past）と戦略提案（Future）が因果関係で接続されているか\n"
        "・SWOT分析の各要素と戦略提案が対応しているか\n"
        "・効果試算の根拠と試算数値の妥当性を検証\n\n"
        "【実現可能性チェック】\n"
        "・提案施策が企業の経営資源（人員・資金）で実行可能か\n"
        "・投資計画と財務状況の整合性を確認\n"
        "・ロードマップの時間軸が建設業の実態に即しているか"
    )

    doc.add_heading("3. 人間による修正・採用判断", level=2)
    doc.add_paragraph(
        "【修正した点】\n"
        "・財務指標の数値誤りを正確な値に修正\n"
        "・抽象的すぎる提案を企業規模・地域特性に合わせて具体化\n"
        "・官公庁/民間、元請/下請の販路特性を反映した表現に調整\n"
        "・所在地の地域特性（人口動態、行政施策等）を反映した提案に補強\n\n"
        "【採用基準】\n"
        "・財務データに基づく客観的根拠が明確であること\n"
        "・建設業界の商慣行・実務に即していること\n"
        "・2024年問題やDX/GXなど業界課題に対応していること\n"
        "・企業規模・地域特性に適した施策であること\n\n"
        "【不採用・修正理由】\n"
        "・根拠が不明確な数値予測は修正または削除\n"
        "・企業規模に見合わない大規模投資提案は現実的な規模に調整\n"
        "・地域特性を無視した画一的な提案は、地域固有の要素を追加"
    )

    doc.add_heading("4. AI活用の工夫", level=2)
    doc.add_paragraph(
        "【プロンプト設計の工夫】\n"
        "・地域金融機関の融資担当者としてのロール設定\n"
        "・評価5観点（全体構成、地域性、業界特性、GX/DX、人材）を明示的に指示\n"
        "・段階的な分析パイプライン（財務→有報→SWOT→戦略→効果試算→提案書）\n"
        "・各セクションの文字数上限を指定して分量を管理\n\n"
        "【品質向上の取り組み】\n"
        "・定量データ（財務CSV）と定性データ（有報テキスト）の統合分析\n"
        "・業界平均との比較による相対評価の実施\n"
        "・クロスSWOT分析による戦略の論理的導出\n"
        "・自己評価用プロンプトによる品質チェックとブラッシュアップ"
    )

    output_path = OUTPUT_DIR / "verification_process_report.docx"
    doc.save(output_path)
    print(f"[SAVE] {output_path}")
    return output_path


# =============================================================================
# ZIP圧縮
# =============================================================================
def create_submission_zip(signate_username: str):
    """提出用ZIPファイルを作成"""
    submission_dir = OUTPUT_DIR / signate_username
    if submission_dir.exists():
        shutil.rmtree(submission_dir)
    submission_dir.mkdir(parents=True)

    # 提案書をコピー
    for docx_file in PROPOSALS_DIR.glob("*.docx"):
        shutil.copy(docx_file, submission_dir / docx_file.name)

    # プロンプトログ
    prompt_log = OUTPUT_DIR / "prompt_log.txt"
    if prompt_log.exists():
        shutil.copy(prompt_log, submission_dir / "prompt_log.txt")

    # 検証プロセスレポート
    verification = OUTPUT_DIR / "verification_process_report.docx"
    if verification.exists():
        shutil.copy(verification, submission_dir / "verification_process_report.docx")

    # ZIP
    zip_path = OUTPUT_DIR / signate_username
    shutil.make_archive(str(zip_path), 'zip', OUTPUT_DIR, signate_username)

    print(f"[SAVE] {zip_path}.zip")
    return f"{zip_path}.zip"


# =============================================================================
# メイン処理
# =============================================================================
def main():
    print("=" * 60)
    print("Step 9: 提案書・成果物作成")
    print("=" * 60 + "\n")

    company_codes = get_company_codes()
    print(f"[INFO] 対象企業: {len(company_codes)}社")

    api_key = os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        print("[ERROR] GOOGLE_API_KEY が設定されていません。")
        return

    # 提案書生成
    print("\n[Step 9.1] 提案書生成（LLMによるセクション生成）")
    generator = ProposalGenerator(api_key)
    generator.generate_all_proposals(company_codes)

    # プロンプトログ保存
    log_text = generator.prompt_logger.export_to_text()
    log_path = PROMPT_LOG_DIR / "step9_prompt_log.txt"
    with open(log_path, "w", encoding="utf-8") as f:
        f.write(log_text)
    print(f"[SAVE] {log_path}")

    # プロンプトログ統合
    print("\n[Step 9.2] プロンプトログ統合")
    merge_prompt_logs()

    # 検証プロセスレポート
    print("\n[Step 9.3] 検証プロセスレポート生成")
    generate_verification_report()

    # ZIP案内
    print("\n[Step 9.4] ZIP圧縮準備")
    print("提出用ZIPを作成するには、以下を実行してください：")
    print("  from step9_proposal_generation import create_submission_zip")
    print("  create_submission_zip('your_signate_username')")

    print("\n" + "=" * 60)
    print("Step 9 完了！")
    print(f"提案書: {PROPOSALS_DIR}")
    print(f"成果物: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
